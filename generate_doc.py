"""生成《时间当铺》游戏策划案 Word 文档 v3.0"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# ============================================================
# 全局样式设置
# ============================================================
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    hs.font.color.rgb = RGBColor(0x2a, 0x1a, 0x3e) if level == 1 else RGBColor(0x4a, 0x3a, 0x6a)

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)


def add_title_page():
    """封面页"""
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('《时间当铺》')
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x2a, 0x1a, 0x3e)
    run.bold = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Time Pawnshop')
    run2.font.size = Pt(18)
    run2.font.color.rgb = RGBColor(0x6a, 0x5a, 0x8a)

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('游 戏 策 划 案')
    run3.font.size = Pt(22)
    run3.font.color.rgb = RGBColor(0x4a, 0x3a, 0x6a)

    for _ in range(4):
        doc.add_paragraph()

    info_lines = [
        '类    型：叙事解谜 / 像素风',
        '平    台：PC / Web / 移动端',
        '开发引擎：HTML5 Canvas + JavaScript',
        '目标人群：18-35岁，喜爱叙事向独立游戏的玩家',
        '预计时长：单次通关 2-3 小时，全结局约 8 小时',
        '文档版本：v3.0',
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()


def add_table(headers, rows, col_widths=None):
    """创建格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_bullet(text, level=0):
    """添加项目符号列表"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 1)
    return p


def add_quote(text):
    """添加引用样式段落"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.5)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x6a, 0x5a, 0x8a)
    run.font.size = Pt(11)
    return p


# ============================================================
# 开始构建文档
# ============================================================

add_title_page()

# ---- 目录页 ----
doc.add_heading('目  录', level=1)
toc_items = [
    '一、项目概述',
    '二、核心玩法设计',
    '    2.1 时间交易系统',
    '    2.2 蝴蝶效应系统',
    '    2.3 共振涟漪系统',
    '    2.4 当铺经营系统',
    '    2.5 记忆潜入系统',
    '三、叙事设计',
    '    3.1 叙事结构',
    '    3.2 角色设定',
    '    3.3 客人故事详案',
    '    3.4 主线剧情',
    '    3.5 结局设计',
    '四、关卡设计',
    '    4.1 游戏流程总览',
    '    4.2 各日关卡配置',
    '    4.3 难度曲线',
    '五、UI / UX 设计',
    '    5.1 界面布局',
    '    5.2 像素美术风格',
    '    5.3 音效与音乐',
    '六、数值设计',
    '    6.1 时间经济系统',
    '    6.2 商店价格体系',
    '    6.3 蝴蝶效应映射表',
    '    6.4 共振涟漪映射表',
    '七、技术方案',
    '八、开发规划',
    '九、竞品分析与市场定位',
    '十、风险评估与应对',
    '十一、成就系统',
    '    11.1 成就分类',
    '    11.2 成就列表',
    '    11.3 成就展示（记忆殿堂）',
    '    11.4 奖励系统',
    '十二、未来发展方向',
    '    12.1 近期优化',
    '    12.2 中期扩展',
    '    12.3 远期规划',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.line_spacing = 1.8
    if not item.startswith('    '):
        for run in p.runs:
            run.bold = True

doc.add_page_break()

# ============================================================
# 一、项目概述
# ============================================================
doc.add_heading('一、项目概述', level=1)

doc.add_heading('1.1 一句话描述', level=2)
add_quote('经营一家收购"时间"的当铺，通过改写客人的过去来改变现在。')

doc.add_heading('1.2 项目愿景', level=2)
doc.add_paragraph(
    '《时间当铺》是一款以"时间"为核心主题的叙事解谜游戏。玩家扮演一位失忆的当铺老板，'
    '在帮助客人改写过去的过程中，逐步揭开自己的身世之谜。游戏通过紧密关联的蝴蝶效应系统，'
    '让每一个选择都真实地影响后续剧情走向，营造出"每一个决定都有重量"的沉浸式体验。'
    'v3.0版本新增三位客人、共振涟漪系统和隐藏结局，构建了完整的六人命运之网。'
)

doc.add_heading('1.3 核心体验支柱', level=2)
pillars = [
    ('选择的重量', '每一个选择都会通过蝴蝶效应系统产生可见的连锁反应，没有"正确答案"。'),
    ('温暖的忧伤', '故事基调是治愈系的忧伤——失去不可逆，但过程中的温暖是真实的。'),
    ('解谜的惊喜', '玩家逐步发现所有客人之间的隐藏联系，以及自己与当铺的真正关系。'),
    ('经营的沉浸', '当铺不仅是叙事容器，也是玩家的"家"——修缮它、升级它，有归属感。'),
    ('命运之网', '六位客人的选择交织成网，共振涟漪在看不见的地方回响。'),
]
for title, desc in pillars:
    p = doc.add_paragraph()
    run_t = p.add_run(f'【{title}】')
    run_t.bold = True
    run_t.font.color.rgb = RGBColor(0x4a, 0x3a, 0x6a)
    p.add_run(f'  {desc}')

doc.add_heading('1.4 产品定位', level=2)
add_table(
    ['维度', '描述'],
    [
        ['类型', '叙事解谜 + 轻度经营'],
        ['风格', '像素风（16色调色板为主），暗色调为基础，暖色点缀'],
        ['视角', '第三人称俯视 / 侧视混合'],
        ['平台', 'PC（Web浏览器）、移动端（响应式适配）'],
        ['时长', '单线通关 2-3 小时，全结局探索约 8 小时'],
        ['定价策略', '免费体验 Demo（第一位客人），完整版付费'],
        ['目标评级', '无暴力/无血腥，全年龄适宜'],
    ],
    col_widths=[4, 12]
)

doc.add_page_break()

# ============================================================
# 二、核心玩法设计
# ============================================================
doc.add_heading('二、核心玩法设计', level=1)

doc.add_heading('2.1 时间交易系统', level=2)
doc.add_paragraph(
    '时间交易是游戏的核心循环。客人带着"记忆片段"来到当铺，将其作为抵押品交给玩家。'
    '玩家进入客人的记忆，在关键时间节点做出选择，从而改变（或维持）该客人的命运。'
    '完成交易后获得"时间碎片"作为报酬。'
)

doc.add_paragraph('交易流程：')
steps = [
    '客人来店 → 对话了解背景故事',
    '接收记忆 → 进入"记忆潜入"场景',
    '关键抉择 → 在记忆节点做出选择',
    '涟漪扩散 → 蝴蝶效应系统计算影响',
    '回到当铺 → 获得时间碎片奖励',
]
for i, s in enumerate(steps):
    add_bullet(f'第{i+1}步：{s}')

doc.add_heading('2.2 蝴蝶效应系统', level=2)
doc.add_paragraph(
    '蝴蝶效应系统是本作最核心的差异化设计。v3.0采用数据驱动架构（BUTTERFLY_EFFECTS表），'
    '支持两种效应类型：block_and_replace（阻止角色出现并播放替代对话）和'
    'greeting_variant（角色仍出现但根据涟漪改变问候语）。'
    '六位客人形成完整的因果之网，每个选择都产生可见的连锁反应。'
)

doc.add_paragraph('蝴蝶效应网络（6人互联）：')
add_table(
    ['源角色', '源选择', '目标角色', '效应类型', '效果'],
    [
        ['陈伯', '学医', '林小姐', '阻止出现', '林小姐母亲被救，不来当铺'],
        ['林小姐', '保守治疗', '小周', '变体问候', '温暖传递，小周不那么害怕'],
        ['林小姐', '手术', '小周', '变体问候', '遗憾回响，小周想起信的内容'],
        ['陈伯', '从军', '阿蘅', '变体问候', '沉默传递，钟楼更加沉重'],
        ['陈伯', '学医', '阿蘅', '变体问候', '智慧传递，钟声格外清脆'],
        ['小周', '寄信', '老吴', '变体问候', '执着传递，老吴想起底片'],
        ['小周', '烧信', '老吴', '变体问候', '犹疑传递，老吴更加纠结'],
    ],
    col_widths=[2, 3, 2, 3, 6]
)

doc.add_heading('2.3 共振涟漪系统', level=2)
doc.add_paragraph(
    '共振涟漪是v3.0新增的机制。当多条特定涟漪同时存在时，它们会产生"共振"，'
    '触发特殊效果和额外的时间碎片奖励。共振涟漪在记忆归还时自动检测。'
)

add_table(
    ['共振名称', '所需涟漪组合', '奖励'],
    [
        ['「医者的信」', '陈伯学医 + 林小姐保守 + 小周寄信', '+25碎片'],
        ['「沉默的钟楼」', '陈伯从军 + 阿蘅停钟 + 老吴销毁', '+25碎片'],
        ['「灰烬中的秒针」', '小周烧信 + 林小姐手术 + 阿蘅走钟', '+30碎片'],
    ],
    col_widths=[4, 7, 3]
)

doc.add_heading('2.4 当铺经营系统', level=2)
doc.add_paragraph(
    '当铺是玩家的"基地"。通过消耗时间碎片来修缮和升级当铺，'
    '解锁新的功能区域、新的客人线索、以及主角自己的记忆碎片。'
    'v3.0新增三项升级，当铺等级从3级扩展至5级。'
)

doc.add_paragraph('升级项目：')
add_table(
    ['升级项', '费用', '所需等级', '授予等级', '效果'],
    [
        ['修复门铃', '20', '0', '1', '解锁第二位客人（林小姐）'],
        ['打磨沙漏', '25', '0', '—', '+50时间碎片，延长时限'],
        ['点亮灯笼', '30', '1', '—', '解锁记忆碎片#1'],
        ['修复铜镜', '40', '1', '2', '解锁更多功能'],
        ['挂起风铃', '45', '2', '3', '解锁更多客人（小周/阿蘅/老吴）'],
        ['点香炉', '35', '2', '—', '额外记忆碎片+20'],
        ['校准钟表', '55', '3', '4', '解锁最终来客（神秘老人）'],
    ],
    col_widths=[3, 2, 2, 2, 7]
)

doc.add_heading('2.5 记忆潜入系统', level=2)
doc.add_paragraph(
    '记忆潜入是游戏的核心玩法场景。玩家"进入"客人的记忆，体验一段特定的时间片段。'
    'v3.0新增三个记忆场景（邮局、钟楼、照相馆），以及隐藏选项注入机制。'
)

doc.add_page_break()

# ============================================================
# 三、叙事设计
# ============================================================
doc.add_heading('三、叙事设计', level=1)

doc.add_heading('3.1 叙事结构', level=2)
doc.add_paragraph('游戏采用三层嵌套叙事结构：')

p = doc.add_paragraph()
run = p.add_run('表面线（单元剧）：')
run.bold = True
p.add_run('每位客人是一个独立的短篇故事。v3.0有六位客人，覆盖不同年代和人生主题。')

p = doc.add_paragraph()
run = p.add_run('连接线（蝴蝶效应+共振涟漪）：')
run.bold = True
p.add_run('六位客人通过蝴蝶效应形成因果之网。特定组合触发共振涟漪，揭示更深层的联系。')

p = doc.add_paragraph()
run = p.add_run('暗线（主线）：')
run.bold = True
p.add_run('六块记忆碎片逐步揭露当铺老板的身世。最终发现所有客人都是自己在不同时间线上的可能性。'
          '集齐特定五个涟漪可解锁隐藏结局「闭环」。')

doc.add_heading('3.2 角色设定', level=2)

# 角色解锁表
doc.add_paragraph('角色解锁进度：')
add_table(
    ['角色', '解锁日', '所需等级', '时间奖励', '记忆场景'],
    [
        ['陈伯', '第1天', '0', '30', '1962年·岔路口'],
        ['林小姐', '第2天', '1', '35', '1985年·省城医院'],
        ['小周', '第3天', '2', '40', '2019年·邮局'],
        ['阿蘅', '第4天', '3', '45', '2005年·钟楼'],
        ['老吴', '第5天', '3', '40', '1998年·照相馆'],
        ['神秘老人', '第6天', '4', '50', '????年·镜中当铺'],
    ],
    col_widths=[3, 2, 2, 2, 5]
)

# 新角色设定
doc.add_heading('小周 — 最后的邮差', level=3)
add_table(
    ['属性', '描述'],
    [
        ['外观', '年轻男性，穿着蓝色邮递员制服，背着褪色的邮包'],
        ['性格', '善良但犹豫不决，内心有强烈的使命感'],
        ['核心冲突', '寄出信件 vs 烧毁信件——"一封信的份量有多重？"'],
        ['记忆时间', '2019年·秋天'],
        ['记忆地点', '县城邮局分拣室'],
        ['蝴蝶效应', '受林小姐选择的涟漪影响（保守→温暖传递/手术→遗憾回响）'],
    ],
    col_widths=[3, 13]
)

doc.add_heading('阿蘅 — 钟楼的守望者', level=3)
add_table(
    ['属性', '描述'],
    [
        ['外观', '沉静的年轻女子，指尖沾着齿轮油，琥珀色调服饰'],
        ['性格', '坚忍，对时间有独特的感悟，话语中带着钟声的节奏'],
        ['核心冲突', '让钟走到尽头 vs 停钟修复——"坚守到最后还是为未来妥协？"'],
        ['记忆时间', '2005年·秋天'],
        ['记忆地点', '县城百年钟楼'],
        ['蝴蝶效应', '受陈伯选择的涟漪影响（从军→沉默传递/学医→智慧传递）'],
    ],
    col_widths=[3, 13]
)

doc.add_heading('老吴 — 消逝的影像', level=3)
add_table(
    ['属性', '描述'],
    [
        ['外观', '老年男性，灰色系服饰，怀抱一台布满裂痕的旧相机'],
        ['性格', '内敛、怀旧，用镜头记录世界却不敢直面自己的记忆'],
        ['核心冲突', '冲洗底片 vs 销毁底片——"该不该让过去重见天日？"'],
        ['记忆时间', '1998年·冬天'],
        ['记忆地点', '县城老街照相馆暗房'],
        ['蝴蝶效应', '受小周选择的涟漪影响（寄信→执着传递/烧信→犹疑传递）'],
    ],
    col_widths=[3, 13]
)

doc.add_heading('3.3 客人故事详案', level=2)

doc.add_heading('故事四：小周 —「未投递的信」', level=3)
doc.add_paragraph(
    '2019年秋天，县城邮局。实习邮递员小周在分拣室发现一封三个月前的信，'
    '收件地址已拆迁。信封未封好，能看到一位父亲写给远方女儿的话，充满思念和歉意。'
)
p = doc.add_paragraph()
run = p.add_run('选择A — 寄出：')
run.bold = True
p.add_run('三个月后信终于送达，但写信的父亲已去世两个月。那封信成了女儿最后的安慰。')
p = doc.add_paragraph()
run = p.add_run('选择B — 烧毁：')
run.bold = True
p.add_run('女儿永远不知道父亲写过这封信。小周也从此离开邮局。')

doc.add_heading('故事五：阿蘅 —「钟楼」', level=3)
doc.add_paragraph(
    '2005年秋天，县城百年钟楼。阿蘅从父亲手中接过守钟人的钥匙。主齿轮严重磨损，'
    '工程师警告：再走两年机芯就会崩溃。现在停下来修复可以再走一百年，但这座城从未停过钟。'
)
p = doc.add_paragraph()
run = p.add_run('选择A — 让钟走到尽头：')
run.bold = True
p.add_run('两年后除夕夜，钟敲完最后一声后齿轮断裂。整座城市的人都停下了脚步。')
p = doc.add_paragraph()
run = p.add_run('选择B — 停钟修复：')
run.bold = True
p.add_run('修复花了一年。钟重启时只有阿蘅一人在场——人们已经习惯了手机看时间。')

doc.add_heading('故事六：老吴 —「暗房」', level=3)
doc.add_paragraph(
    '1998年冬天，县城老街照相馆暗房。老吴找出一卷二十年前的底片，'
    '上面记录着他和初恋的最后一个秋天。她第二年就去了远方，再也没有回来。'
)
p = doc.add_paragraph()
run = p.add_run('选择A — 冲洗底片：')
run.bold = True
p.add_run('照片放在橱窗里，三年后初恋的女儿找上门来，断裂的缘分重新接上。')
p = doc.add_paragraph()
run = p.add_run('选择B — 销毁底片：')
run.bold = True
p.add_run('老吴关了暗房，不再冲洗胶卷。但每年秋天，他都会走一走她们曾经散步的路。')

doc.add_heading('3.4 主线剧情', level=2)
doc.add_paragraph(
    '主线通过六块"老板记忆碎片"渐进式揭露（v3.0从3块扩展至6块）：'
)

add_table(
    ['碎片', '解锁条件', '内容'],
    [
        ['碎片一：雨中的身影', '购买"点亮灯笼"', '一个雨天，有人在等你。'],
        ['碎片二：破碎的承诺', '完成陈伯故事', '"我会回来的。"你对谁说过？'],
        ['碎片三：时间的代价', '完成林小姐故事', '你用全部过去换取了"再来一次"。'],
        ['碎片四：信上的字迹', '完成小周故事', '一封自己写的信："请替我活下去。"'],
        ['碎片五：停摆的钟', '完成阿蘅故事', '你内心的时间第一次停止流动。'],
        ['碎片六：最后的影像', '完成老吴故事', '你和一个女人的合影。你记起了她的名字。'],
    ],
    col_widths=[4, 4, 8]
)

doc.add_heading('3.5 结局设计', level=2)
doc.add_paragraph('游戏共有四种结局（v3.0新增隐藏结局D）：')

p = doc.add_paragraph()
run = p.add_run('结局A — 晨光（关闭当铺）')
run.bold = True
doc.add_paragraph('玩家选择打破时间循环。基调：释然、温暖。')

p = doc.add_paragraph()
run = p.add_run('结局B — 永恒（继续经营）')
run.bold = True
doc.add_paragraph('玩家选择成为永恒的时间守护者。基调：豁达、使命感。')

p = doc.add_paragraph()
run = p.add_run('结局C — 时间耗尽（默认结局）')
run.bold = True
doc.add_paragraph('时间碎片耗尽。基调：遗憾、鼓励重玩。')

p = doc.add_paragraph()
run = p.add_run('结局D — 闭环（隐藏结局）')
run.bold = True
run.font.color.rgb = RGBColor(0xaa, 0x88, 0xff)
doc.add_paragraph(
    '触发条件：在神秘老人的记忆中，集齐五个特定涟漪'
    '（陈伯学医、林小姐保守、小周烧信、阿蘅走钟、老吴冲洗）后出现第三个选项。'
    '五条时间线交汇，当铺变成时间之网的节点。玩家不是囚徒也不是掌柜，而是时间编织者。'
    '基调：超越、顿悟。'
)

doc.add_page_break()

# ============================================================
# 四、关卡设计
# ============================================================
doc.add_heading('四、关卡设计', level=1)

doc.add_heading('4.1 游戏流程总览', level=2)

add_table(
    ['天数', '可用事件', '可用升级', '剧情推进'],
    [
        ['第1天', '陈伯来访', '—', '开场，第一段记忆'],
        ['第2天', '林小姐来访/蝴蝶效应', '修复门铃(20)、沙漏(25)', '蝴蝶效应初触发'],
        ['第3天', '小周来访', '铜镜(40)、灯笼(30)', '第三段记忆'],
        ['第4天', '阿蘅来访', '风铃(45)、香炉(35)', '钟楼记忆'],
        ['第5天', '老吴来访', '—', '照相馆记忆'],
        ['第6天+', '神秘老人来访', '钟表(55)', '真相揭示、最终选择'],
    ],
    col_widths=[3, 5, 5, 5]
)

doc.add_page_break()

# ============================================================
# 六、数值设计
# ============================================================
doc.add_heading('六、数值设计', level=1)

doc.add_heading('6.1 时间经济系统', level=2)

add_table(
    ['参数', '数值', '说明'],
    [
        ['初始时间碎片', '100', '足够支撑约20天基础消耗'],
        ['每日消耗', '5', '即使不操作也消耗'],
        ['陈伯奖励', '30', '第一位客人'],
        ['林小姐奖励', '35', '第二位客人'],
        ['小周奖励', '40', '第三位客人'],
        ['阿蘅奖励', '45', '第四位客人'],
        ['老吴奖励', '40', '第五位客人'],
        ['神秘老人奖励', '50', '最终客人'],
        ['蝴蝶效应奖励', '15', '阻止型蝴蝶效应奖励'],
        ['共振涟漪奖励', '25-30', '三种共振涟漪'],
    ],
    col_widths=[4, 3, 9]
)

doc.add_heading('6.3 蝴蝶效应映射表', level=2)
add_table(
    ['源选择', '涟漪ID', '影响目标', '效果类型'],
    [
        ['陈伯→学医', 'chen_doctor', '林小姐', '阻止出现'],
        ['陈伯→从军', 'chen_army', '阿蘅', '变体问候'],
        ['陈伯→学医', 'chen_doctor', '阿蘅', '变体问候'],
        ['林小姐→保守', 'lin_conservative', '小周', '变体问候'],
        ['林小姐→手术', 'lin_surgery', '小周', '变体问候'],
        ['小周→寄信', 'zhou_send', '老吴', '变体问候'],
        ['小周→烧信', 'zhou_burn', '老吴', '变体问候'],
    ],
    col_widths=[4, 3, 3, 6]
)

doc.add_heading('6.4 共振涟漪映射表', level=2)
add_table(
    ['共振名称', '涟漪1', '涟漪2', '涟漪3', '奖励'],
    [
        ['医者的信', 'chen_doctor', 'lin_conservative', 'zhou_send', '+25'],
        ['沉默的钟楼', 'chen_army', 'heng_clock_stop', 'wu_destroy', '+25'],
        ['灰烬中的秒针', 'zhou_burn', 'lin_surgery', 'heng_clock_run', '+30'],
    ],
    col_widths=[4, 3, 3, 3, 2]
)

doc.add_page_break()

# ============================================================
# 七、技术方案
# ============================================================
doc.add_heading('七、技术方案', level=1)

doc.add_heading('7.1 架构设计', level=2)
doc.add_paragraph('代码采用模块化架构，分为三层：')

p = doc.add_paragraph()
run = p.add_run('数据层 (data.js, ~2000行)：')
run.bold = True
add_bullet('SPRITES — 10个像素精灵定义')
add_bullet('CLIENTS — 6位客人数据')
add_bullet('STORIES — 6组故事/对话树 + 彩蛋对话')
add_bullet('BACKGROUNDS — 7个程序化背景绘制函数')
add_bullet('BUTTERFLY_EFFECTS — 蝴蝶效应数据表（数据驱动）')
add_bullet('RESONANCE_EFFECTS — 共振涟漪数据表')
add_bullet('SHOP_ITEMS — 7项商店升级')
add_bullet('OWNER_MEMORIES — 6块主线记忆碎片')
add_bullet('ACHIEVEMENTS — 40+项成就定义')
add_bullet('ALL_RIPPLE_IDS — 13个涟漪ID索引')

p = doc.add_paragraph()
run = p.add_run('逻辑层 (game.js, ~1950行)：')
run.bold = True
add_bullet('GameState — 含数据驱动的_isClientBlocked()和triggeredResonances')
add_bullet('ShopScene — 含_selectGreeting()、_checkResonanceEffects()、7节点涟漪图谱')
add_bullet('MemoryScene — 含_buildChoices()隐藏选项注入')
add_bullet('EndingScene — 含ending_loop分支')
add_bullet('AchievementSystem — 40+成就检测')

p = doc.add_paragraph()
run = p.add_run('引擎层 (engine.js, ~850行)：')
run.bold = True
add_bullet('AudioManager — 含resonance和弦音效')
add_bullet('其余引擎模块保持不变')

doc.add_page_break()

# ============================================================
# 十一、成就系统
# ============================================================
doc.add_heading('十一、成就系统', level=1)

doc.add_paragraph(
    'v3.0成就系统扩展至40+项，新增抉择之重（10项覆盖6位客人选择）、'
    '共振相关成就、隐藏结局成就「闭环」和「四重门」等。'
)

doc.add_heading('11.1 新增成就', level=2)
add_table(
    ['成就', '条件', '分类', '稀有度'],
    [
        ['使命必达', '帮助小周选择寄出信件', '抉择之重', '★★'],
        ['付之一炬', '帮助小周选择烧毁信件', '抉择之重', '★★'],
        ['最后的钟声', '帮助阿蘅选择让钟继续走', '抉择之重', '★★'],
        ['沉默之声', '帮助阿蘅选择让钟停下来', '抉择之重', '★★'],
        ['重见天日', '帮助老吴选择冲洗底片', '抉择之重', '★★'],
        ['心中的影像', '帮助老吴选择销毁底片', '抉择之重', '★★'],
        ['共鸣', '触发第一次共振涟漪', '蝴蝶之翼', '★★★'],
        ['交响曲', '触发全部共振涟漪', '蝴蝶之翼', '★★★★★'],
        ['六度分隔', '完成全部六位客人的故事', '蝴蝶之翼', '★★★★'],
        ['闭环', '发现隐藏结局D', '结局收藏(隐)', '★★★★★'],
        ['四重门', '见证全部四种结局', '结局收藏(隐)', '★★★★★'],
        ['碎片·字迹', '发现第四块记忆碎片', '记忆深处', '★★'],
        ['碎片·停摆', '发现第五块记忆碎片', '记忆深处', '★★'],
        ['碎片·影像', '发现第六块记忆碎片', '记忆深处', '★★★'],
    ],
    col_widths=[3, 5, 3, 2]
)

doc.add_page_break()

# ============================================================
# 十二、未来发展方向
# ============================================================
doc.add_heading('十二、未来发展方向', level=1)

doc.add_heading('12.1 v3.0 完成内容', level=2)
add_bullet('六位客人完整故事线（陈伯、林小姐、小周、阿蘅、老吴、神秘老人）')
add_bullet('数据驱动蝴蝶效应系统（BUTTERFLY_EFFECTS表）')
add_bullet('共振涟漪系统（3组共振效果）')
add_bullet('隐藏结局D「闭环」（五涟漪条件触发）')
add_bullet('7节点涟漪图谱可视化')
add_bullet('5级当铺等级 + 7项商店升级')
add_bullet('6块老板记忆碎片')
add_bullet('40+项成就')

doc.add_heading('12.2 后续规划', level=2)
add_bullet('存档书签系统：支持手动存档和读档')
add_bullet('画廊模式：收集已解锁的精灵、CG、音乐')
add_bullet('多语言支持：英语、日语等')
add_bullet('Steam发行：上架Steam平台')

doc.add_paragraph()
doc.add_paragraph()

# 结语
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 文档结束 —')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x6a, 0x5a, 0x8a)
run.italic = True

# ============================================================
# 保存文档
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), '《时间当铺》游戏策划案.docx')
doc.save(output_path)
print(f'文档已生成: {output_path}')
